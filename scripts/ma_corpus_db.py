"""Centralized training corpus database and ingestion utilities."""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import sqlite3
import threading
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SQLITE_PATH = ROOT / "instance" / "lawagent_training.sqlite3"
DEPOSIT_DIRS = [
    ROOT / "attached_assets",
    ROOT / "training_docs_inbox",
    ROOT / "training_docs_inbox" / "guides",
    ROOT / "training_docs_inbox" / "notes",
    ROOT / "training_docs_inbox" / "playbooks",
]
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class CorpusDocument:
    id: int
    title: str
    source_path: str
    category: str
    document_type: str
    source_system: str
    checksum: str
    chunk_count: int
    created_at: str
    jurisdiction: str = ""
    deal_stance: str = ""
    deal_structure: str = ""


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def file_checksum(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def extract_text(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        docs = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(page_content=text, metadata={"page": page_index}))
        return docs
    if suffix == ".docx":
        doc = DocxDocument(path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return [Document(page_content=text, metadata={"page": 1})] if text.strip() else []
    if suffix in {".txt", ".md"}:
        return [Document(page_content=path.read_text(encoding="utf-8", errors="ignore"), metadata={"page": 1})]
    return []


FOLDER_CATEGORY_HINTS: dict[str, tuple[str, str]] = {
    "guides": ("guide", "Guide"),
    "notes": ("capability_notes", "Capability Notes"),
    "playbooks": ("playbook", "Playbook"),
}


def classify_document(title: str, text: str, folder_hint: str = "") -> dict[str, str]:
    haystack = f"{title}\n{text[:8000]}".lower()
    rules = [
        ("ancillary_agreements", "Resource Kit", ["ancillary", "escrow agreements", "transition services", "assignment and assumption"]),
        ("asset_acquisition", "Resource Kit", ["asset acquisition", "asset purchase", "bill of sale", "assumed liabilities"]),
        ("due_diligence", "Resource Kit", ["due diligence", "request list", "diligence checklist", "specialist areas"]),
        ("corporate_templates_market_data", "Template Resource", ["market data", "templates integrated", "practical guidance content type"]),
        ("ip_technology", "Specialist Diligence", ["intellectual property", "software", "open source", "cybersecurity", "privacy"]),
        ("employment_benefits", "Specialist Diligence", ["employment", "employee benefits", "executive compensation", "change-in-control"]),
        ("regulatory", "Specialist Diligence", ["antitrust", "fcpa", "regulatory", "sanctions", "export controls"]),
        ("environmental", "Specialist Diligence", ["environmental", "climate change", "esg"]),
        ("real_estate", "Specialist Diligence", ["real estate", "reit", "property"]),
        ("purchase_agreement", "Agreement Template", ["agreement and plan", "purchase agreement", "merger agreement", "representations and warranties"]),
        ("guide", "Guide", ["guide", "how to", "step-by-step", "walkthrough", "handbook", "getting started", "best practice"]),
        ("practical_guidance", "Practical Guidance", ["practical guidance", "practice note", "practice point", "drafting note", "advisory", "key considerations"]),
        ("playbook", "Playbook", ["playbook", "runbook", "workflow", "standard operating procedure", "sop", "action plan", "execution plan"]),
        ("capability_notes", "Capability Notes", ["capability", "agent note", "improvement note", "capability improvement", "skill assessment", "performance note"]),
        ("prompt_engineering", "Prompt Engineering", ["prompt engineering", "prompt template", "system prompt", "few-shot", "chain of thought", "prompt design", "llm instruction"]),
        ("training_instructions", "Training Instructions", ["training instruction", "training material", "onboarding", "curriculum", "learning objective", "training guide", "lesson plan"]),
    ]
    for category, doc_type, keywords in rules:
        if any(keyword in haystack for keyword in keywords):
            return {"category": category, "document_type": doc_type}

    if folder_hint and folder_hint in FOLDER_CATEGORY_HINTS:
        cat, dtype = FOLDER_CATEGORY_HINTS[folder_hint]
        return {"category": cat, "document_type": dtype}

    return {"category": "general_ma", "document_type": "Training Document"}


def detect_tags(title: str, text: str) -> dict[str, str]:
    haystack = f"{title}\n{text[:12000]}".lower()

    jurisdiction = ""
    jurisdiction_rules = [
        ("Delaware", ["delaware", "del. code", "court of chancery", "8 del. c."]),
        ("New York", ["new york", "n.y.", "nyc", "manhattan"]),
        ("California", ["california", "cal. corp. code", "calif."]),
        ("Texas", ["texas", "tex. bus. org. code"]),
        ("Nevada", ["nevada", "nev. rev. stat"]),
        ("Illinois", ["illinois", "ill. bus. corp. act"]),
        ("United Kingdom", ["english law", "uk", "united kingdom", "companies act 2006"]),
        ("Canada", ["canada", "canadian", "cbca", "ontario"]),
        ("Federal/Multi-State", ["federal", "multi-state", "multiple jurisdictions"]),
    ]
    for label, keywords in jurisdiction_rules:
        if any(kw in haystack for kw in keywords):
            jurisdiction = label
            break

    deal_stance = ""
    pro_seller_signals = [
        "seller-friendly", "pro-seller", "seller favorable", "seller's sole discretion",
        "no survival", "no indemnification", "as-is", "limited representations",
        "seller shall not be liable", "cap on liability", "de minimis",
    ]
    pro_buyer_signals = [
        "buyer-friendly", "pro-buyer", "buyer favorable", "buyer's sole discretion",
        "full indemnification", "unlimited survival", "bring-down condition",
        "specific performance", "material adverse effect", "buyer's option",
        "escrow fund", "holdback amount", "extensive representations",
    ]
    seller_hits = sum(1 for s in pro_seller_signals if s in haystack)
    buyer_hits = sum(1 for s in pro_buyer_signals if s in haystack)
    if seller_hits > buyer_hits and seller_hits >= 2:
        deal_stance = "pro-seller"
    elif buyer_hits > seller_hits and buyer_hits >= 2:
        deal_stance = "pro-buyer"
    elif seller_hits >= 1 or buyer_hits >= 1:
        deal_stance = "balanced"

    deal_structure = ""
    if any(phrase in haystack for phrase in ["asset purchase", "asset acquisition", "bill of sale", "assumed liabilities", "purchased assets"]):
        deal_structure = "asset purchase"
    elif any(phrase in haystack for phrase in ["stock purchase", "equity purchase", "share purchase", "all outstanding shares", "stock acquisition"]):
        deal_structure = "stock purchase"
    elif any(phrase in haystack for phrase in ["merger agreement", "agreement and plan of merger", "surviving corporation", "merger sub"]):
        deal_structure = "merger"

    return {
        "jurisdiction": jurisdiction,
        "deal_stance": deal_stance,
        "deal_structure": deal_structure,
    }


def detect_source_system(text: str, path: Path | None = None) -> str:
    lower = text[:3000].lower()
    if "lexisnexis" in lower or "practical guidance" in lower:
        return "LexisNexis Practical Guidance user-provided export"
    if "sec.gov" in lower or "edgar" in lower or "securities and exchange commission" in lower:
        return "SEC EDGAR public filing"
    if path and "edgar" in str(path).lower():
        return "SEC EDGAR public filing"
    if "securities exchange act" in lower or "form 8-k" in lower or "form 10-k" in lower:
        return "SEC EDGAR public filing"
    return "User-provided local document"


def tokenize(text: str) -> set[str]:
    return {token for token in re.findall(r"[a-zA-Z][a-zA-Z0-9_]{2,}", text.lower())}


class CorpusDatabase:
    _schema_initialized = False
    _schema_lock = threading.Lock()

    def __init__(self) -> None:
        self.database_url = os.environ.get("DATABASE_URL")
        self.sqlite_path = Path(os.environ.get("LAWAGENT_SQLITE_PATH", DEFAULT_SQLITE_PATH))
        self.backend = "postgres" if self.database_url else "sqlite"

    def connect(self):
        if self.backend == "postgres":
            import psycopg
            from psycopg.rows import dict_row

            return psycopg.connect(self.database_url, row_factory=dict_row, connect_timeout=10)
        self.sqlite_path.parent.mkdir(parents=True, exist_ok=True)
        connection = sqlite3.connect(self.sqlite_path, timeout=10)
        connection.row_factory = sqlite3.Row
        return connection

    def init_schema(self) -> None:
        if CorpusDatabase._schema_initialized:
            return
        with CorpusDatabase._schema_lock:
            if CorpusDatabase._schema_initialized:
                return
            self._create_tables()

    def _create_tables(self) -> None:
        with self.connect() as connection:
            if self.backend == "postgres":
                connection.execute(
                    """
                    CREATE TABLE IF NOT EXISTS lawagent_documents (
                        id SERIAL PRIMARY KEY,
                        title TEXT NOT NULL,
                        source_path TEXT UNIQUE NOT NULL,
                        category TEXT NOT NULL,
                        document_type TEXT NOT NULL,
                        source_system TEXT NOT NULL,
                        checksum TEXT NOT NULL,
                        metadata JSONB NOT NULL DEFAULT '{}'::jsonb,
                        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                    )
                    """
                )
                connection.execute(
                    """
                    CREATE TABLE IF NOT EXISTS lawagent_chunks (
                        id SERIAL PRIMARY KEY,
                        document_id INTEGER NOT NULL REFERENCES lawagent_documents(id) ON DELETE CASCADE,
                        chunk_index INTEGER NOT NULL,
                        page INTEGER NOT NULL,
                        text TEXT NOT NULL,
                        keywords TEXT NOT NULL,
                        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                    )
                    """
                )
                connection.execute("CREATE INDEX IF NOT EXISTS idx_lawagent_chunks_document_id ON lawagent_chunks(document_id)")
                connection.execute("CREATE INDEX IF NOT EXISTS idx_lawagent_documents_category ON lawagent_documents(category)")
                connection.execute("""
                    ALTER TABLE lawagent_chunks ADD COLUMN IF NOT EXISTS keywords_tsv tsvector
                """)
                connection.execute("""
                    CREATE INDEX IF NOT EXISTS idx_lawagent_chunks_keywords_tsv
                    ON lawagent_chunks USING GIN(keywords_tsv)
                """)
                connection.execute("""
                    UPDATE lawagent_chunks SET keywords_tsv = to_tsvector('english', keywords)
                    WHERE keywords_tsv IS NULL
                """)
            else:
                connection.execute(
                    """
                    CREATE TABLE IF NOT EXISTS lawagent_documents (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        title TEXT NOT NULL,
                        source_path TEXT UNIQUE NOT NULL,
                        category TEXT NOT NULL,
                        document_type TEXT NOT NULL,
                        source_system TEXT NOT NULL,
                        checksum TEXT NOT NULL,
                        metadata TEXT NOT NULL DEFAULT '{}',
                        created_at TEXT NOT NULL
                    )
                    """
                )
                connection.execute(
                    """
                    CREATE TABLE IF NOT EXISTS lawagent_chunks (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        document_id INTEGER NOT NULL,
                        chunk_index INTEGER NOT NULL,
                        page INTEGER NOT NULL,
                        text TEXT NOT NULL,
                        keywords TEXT NOT NULL,
                        created_at TEXT NOT NULL,
                        FOREIGN KEY(document_id) REFERENCES lawagent_documents(id) ON DELETE CASCADE
                    )
                    """
                )
            connection.commit()
        CorpusDatabase._schema_initialized = True

    def upsert_document(self, path: Path, tag_overrides: dict[str, str] | None = None, folder_hint: str = "") -> dict[str, Any]:
        self.init_schema()
        path = path.resolve()
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            return {"status": "skipped", "reason": "unsupported_file", "path": str(path)}

        try:
            checksum = file_checksum(path)
            extracted = extract_text(path)
            full_text = "\n\n".join(doc.page_content for doc in extracted)
        except Exception as exc:
            logger.warning("Failed to parse document %s: %s", path, exc, exc_info=True)
            return {
                "status": "error",
                "reason": "parse_failed",
                "path": str(path),
                "title": path.stem,
                "error": str(exc),
            }

        if not full_text.strip():
            return {"status": "skipped", "reason": "no_text_extracted", "path": str(path), "title": path.stem}

        classification = classify_document(path.stem, full_text, folder_hint=folder_hint)
        source_system = detect_source_system(full_text, path)
        tags = detect_tags(path.stem, full_text)
        if tag_overrides:
            for key in ("jurisdiction", "deal_stance", "deal_structure"):
                if tag_overrides.get(key):
                    tags[key] = tag_overrides[key]
        metadata = {
            "original_filename": path.name,
            "extension": path.suffix.lower(),
            "word_count": len(re.findall(r"\w+", full_text)),
            "ingested_at": now_iso(),
            "jurisdiction": tags["jurisdiction"],
            "deal_stance": tags["deal_stance"],
            "deal_structure": tags["deal_structure"],
        }

        splitter = RecursiveCharacterTextSplitter(chunk_size=1400, chunk_overlap=180)
        chunks = []
        for doc in extracted:
            for split in splitter.split_documents([doc]):
                text = normalize_ws(split.page_content)
                if len(text) >= 80:
                    chunks.append({"text": text, "page": int(split.metadata.get("page", 1) or 1)})

        try:
            with self.connect() as connection:
                existing = self._fetch_one(connection, "SELECT id, checksum, metadata FROM lawagent_documents WHERE source_path = %s", (str(path),))
                if existing and existing["checksum"] == checksum:
                    old_meta = existing.get("metadata") or {}
                    if isinstance(old_meta, str):
                        try:
                            old_meta = json.loads(old_meta)
                        except (json.JSONDecodeError, TypeError):
                            old_meta = {}
                    tags_changed = any(old_meta.get(k) != metadata.get(k) for k in ("jurisdiction", "deal_stance", "deal_structure"))
                    if tags_changed:
                        if self.backend == "postgres":
                            self._execute(connection, "UPDATE lawagent_documents SET metadata = %s::jsonb WHERE id = %s", (json.dumps(metadata), existing["id"]))
                        else:
                            self._execute(connection, "UPDATE lawagent_documents SET metadata = %s WHERE id = %s", (json.dumps(metadata), existing["id"]))
                        connection.commit()
                        return {
                            "status": "tags_updated", "document_id": existing["id"], "path": str(path),
                            "title": path.stem, "category": classification["category"],
                            "document_type": classification["document_type"], "source_system": source_system,
                            "chunk_count": 0, "jurisdiction": tags["jurisdiction"],
                            "deal_stance": tags["deal_stance"], "deal_structure": tags["deal_structure"],
                        }
                    return {"status": "unchanged", "document_id": existing["id"], "path": str(path)}
                if existing:
                    document_id = existing["id"]
                    self._execute(
                        connection,
                        """
                        UPDATE lawagent_documents
                        SET title = %s, category = %s, document_type = %s, source_system = %s, checksum = %s, metadata = %s::jsonb
                        WHERE id = %s
                        """,
                        (
                            path.stem,
                            classification["category"],
                            classification["document_type"],
                            source_system,
                            checksum,
                            json.dumps(metadata),
                            document_id,
                        ),
                    )
                    self._execute(connection, "DELETE FROM lawagent_chunks WHERE document_id = %s", (document_id,))
                else:
                    document_id = self._insert_document(
                        connection,
                        path.stem,
                        str(path),
                        classification["category"],
                        classification["document_type"],
                        source_system,
                        checksum,
                        metadata,
                    )

                for index, chunk in enumerate(chunks):
                    keywords = " ".join(sorted(tokenize(chunk["text"]))[:120])
                    if self.backend == "postgres":
                        self._execute(
                            connection,
                            """
                            INSERT INTO lawagent_chunks (document_id, chunk_index, page, text, keywords, keywords_tsv, created_at)
                            VALUES (%s, %s, %s, %s, %s, to_tsvector('english', %s), %s)
                            """,
                            (document_id, index, chunk["page"], chunk["text"], keywords, keywords, now_iso()),
                        )
                    else:
                        self._execute(
                            connection,
                            """
                            INSERT INTO lawagent_chunks (document_id, chunk_index, page, text, keywords, created_at)
                            VALUES (%s, %s, %s, %s, %s, %s)
                            """,
                            (document_id, index, chunk["page"], chunk["text"], keywords, now_iso()),
                        )
                connection.commit()
        except Exception as exc:
            logger.warning("Failed to upsert document %s: %s", path, exc, exc_info=True)
            return {
                "status": "error",
                "reason": "database_write_failed",
                "path": str(path),
                "title": path.stem,
                "error": str(exc),
            }

        return {
            "status": "ingested" if not existing else "updated",
            "document_id": document_id,
            "title": path.stem,
            "category": classification["category"],
            "document_type": classification["document_type"],
            "source_system": source_system,
            "chunk_count": len(chunks),
            "path": str(path),
            "jurisdiction": tags["jurisdiction"],
            "deal_stance": tags["deal_stance"],
            "deal_structure": tags["deal_structure"],
        }

    @staticmethod
    def _derive_folder_hint(file_path: Path) -> str:
        inbox_root = ROOT / "training_docs_inbox"
        try:
            rel = file_path.resolve().relative_to(inbox_root.resolve())
        except ValueError:
            return ""
        for part in rel.parts:
            if part in FOLDER_CATEGORY_HINTS:
                return part
        return ""

    def ingest_deposit_dirs(self) -> list[dict[str, Any]]:
        for directory in DEPOSIT_DIRS:
            directory.mkdir(parents=True, exist_ok=True)
        seen: set[Path] = set()
        results = []
        for directory in DEPOSIT_DIRS:
            if not directory.exists():
                continue
            for path in sorted(directory.rglob("*")):
                resolved = path.resolve()
                if resolved in seen:
                    continue
                seen.add(resolved)
                if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                    hint = self._derive_folder_hint(path)
                    try:
                        results.append(self.upsert_document(path, folder_hint=hint))
                    except Exception as exc:
                        logger.warning("Deposit ingestion failed for %s: %s", path, exc, exc_info=True)
                        results.append({
                            "status": "error",
                            "reason": "ingest_exception",
                            "path": str(path),
                            "title": path.stem,
                            "error": str(exc),
                        })
        return results

    def delete_document(self, document_id: int) -> dict[str, Any]:
        self.init_schema()
        with self.connect() as connection:
            existing = self._fetch_one(connection, "SELECT id, title FROM lawagent_documents WHERE id = %s", (document_id,))
            if not existing:
                return {"error": "Document not found", "document_id": document_id}
            self._execute(connection, "DELETE FROM lawagent_chunks WHERE document_id = %s", (document_id,))
            self._execute(connection, "DELETE FROM lawagent_documents WHERE id = %s", (document_id,))
            connection.commit()
        return {"status": "deleted", "document_id": document_id, "title": existing["title"]}

    def retrieve(self, query: str, top_k: int = 8, category: str | None = None) -> list[dict[str, Any]]:
        self.init_schema()
        query_tokens = tokenize(query)
        if not query_tokens:
            return []

        select_cols = """
            c.id, c.document_id, c.chunk_index, c.page, c.text, c.keywords,
            d.title, d.category, d.document_type, d.source_system, d.source_path
        """

        if self.backend == "postgres":
            tsquery = " | ".join(sorted(query_tokens)[:20])
            cat_clause = "AND d.category = %s" if category else ""
            sql = f"""
                SELECT {select_cols},
                       ts_rank(c.keywords_tsv, to_tsquery('english', %s)) AS ts_score
                FROM lawagent_chunks c
                JOIN lawagent_documents d ON d.id = c.document_id
                WHERE c.keywords_tsv @@ to_tsquery('english', %s) {cat_clause}
                ORDER BY ts_score DESC
                LIMIT %s
            """
            params: tuple = (tsquery, tsquery)
            if category:
                params += (category,)
            params += (top_k * 3,)
        else:
            keyword_terms = sorted(query_tokens)[:20]
            keyword_clauses = " OR ".join("c.keywords LIKE %s" for _ in keyword_terms)
            keyword_params = tuple(f"%{term}%" for term in keyword_terms)
            cat_clause = "AND d.category = %s" if category else ""
            sql = f"""
                SELECT {select_cols}
                FROM lawagent_chunks c
                JOIN lawagent_documents d ON d.id = c.document_id
                WHERE ({keyword_clauses}) {cat_clause}
            """
            params = keyword_params
            if category:
                params += (category,)

        with self.connect() as connection:
            rows = self._fetch_all(connection, sql, params)

        scored = []
        for row in rows:
            text_tokens = tokenize(f"{row['title']} {row['category']} {row['keywords']} {row['text'][:2500]}")
            overlap = query_tokens & text_tokens
            score = len(overlap)
            if row["category"].replace("_", " ") in query.lower():
                score += 4
            if score:
                scored.append((score, row, overlap))
        ranked = sorted(scored, key=lambda item: item[0], reverse=True)[:top_k]
        return [
            {
                "score": score,
                "document_id": row["document_id"],
                "chunk_id": row["id"],
                "title": row["title"],
                "category": row["category"],
                "document_type": row["document_type"],
                "source_system": row["source_system"],
                "source_path": row["source_path"],
                "page": row["page"],
                "matched_terms": sorted(overlap)[:20],
                "text": row["text"],
            }
            for score, row, overlap in ranked
        ]

    def list_documents(self) -> list[CorpusDocument]:
        self.init_schema()
        with self.connect() as connection:
            rows = self._fetch_all(
                connection,
                """
                SELECT d.id, d.title, d.source_path, d.category, d.document_type, d.source_system,
                       d.checksum, d.metadata, d.created_at, COUNT(c.id) AS chunk_count
                FROM lawagent_documents d
                LEFT JOIN lawagent_chunks c ON c.document_id = d.id
                GROUP BY d.id, d.title, d.source_path, d.category, d.document_type, d.source_system, d.checksum, d.metadata, d.created_at
                ORDER BY d.created_at DESC
                """,
                (),
            )
        results = []
        for row in rows:
            meta = row.get("metadata") or {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except (json.JSONDecodeError, TypeError):
                    meta = {}
            results.append(CorpusDocument(
                id=int(row["id"]),
                title=row["title"],
                source_path=row["source_path"],
                category=row["category"],
                document_type=row["document_type"],
                source_system=row["source_system"],
                checksum=row["checksum"],
                chunk_count=int(row["chunk_count"]),
                created_at=str(row["created_at"]),
                jurisdiction=meta.get("jurisdiction", ""),
                deal_stance=meta.get("deal_stance", ""),
                deal_structure=meta.get("deal_structure", ""),
            ))
        return results

    def get_all_chunks(self) -> list[dict[str, Any]]:
        self.init_schema()
        with self.connect() as connection:
            rows = self._fetch_all(
                connection,
                """
                SELECT c.id, c.document_id, c.page, c.text,
                       d.title, d.category, d.document_type, d.source_system, d.metadata
                FROM lawagent_chunks c
                JOIN lawagent_documents d ON d.id = c.document_id
                ORDER BY c.document_id, c.chunk_index
                """,
                (),
            )
        chunks: list[dict[str, Any]] = []
        for row in rows:
            meta = row.get("metadata") or {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except (json.JSONDecodeError, TypeError):
                    meta = {}
            chunks.append({
                "chunk_id": row["id"],
                "document_id": row["document_id"],
                "page": row["page"],
                "text": row["text"],
                "title": row["title"],
                "category": row["category"],
                "document_type": row["document_type"],
                "source_system": row["source_system"],
                "jurisdiction": meta.get("jurisdiction", ""),
                "deal_stance": meta.get("deal_stance", ""),
                "deal_structure": meta.get("deal_structure", ""),
            })
        return chunks

    def get_chunks_for_documents(self, document_ids: list[int]) -> list[dict[str, Any]]:
        self.init_schema()
        if not document_ids:
            return []

        unique_ids = sorted({int(doc_id) for doc_id in document_ids if isinstance(doc_id, int) or str(doc_id).isdigit()})
        if not unique_ids:
            return []

        placeholders = ", ".join(["%s"] * len(unique_ids))
        sql = f"""
            SELECT c.id, c.document_id, c.page, c.text,
                   d.title, d.category, d.document_type, d.source_system, d.metadata
            FROM lawagent_chunks c
            JOIN lawagent_documents d ON d.id = c.document_id
            WHERE c.document_id IN ({placeholders})
            ORDER BY c.document_id, c.chunk_index
        """
        with self.connect() as connection:
            rows = self._fetch_all(connection, sql, tuple(unique_ids))

        chunks: list[dict[str, Any]] = []
        for row in rows:
            meta = row.get("metadata") or {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except (json.JSONDecodeError, TypeError):
                    meta = {}
            chunks.append({
                "chunk_id": row["id"],
                "document_id": row["document_id"],
                "page": row["page"],
                "text": row["text"],
                "title": row["title"],
                "category": row["category"],
                "document_type": row["document_type"],
                "source_system": row["source_system"],
                "jurisdiction": meta.get("jurisdiction", ""),
                "deal_stance": meta.get("deal_stance", ""),
                "deal_structure": meta.get("deal_structure", ""),
            })
        return chunks

    def stats(self) -> dict[str, Any]:
        documents = self.list_documents()
        categories: dict[str, int] = {}
        total_chunks = 0
        for doc in documents:
            categories[doc.category] = categories.get(doc.category, 0) + 1
            total_chunks += doc.chunk_count
        return {
            "backend": self.backend,
            "document_count": len(documents),
            "chunk_count": total_chunks,
            "categories": categories,
            "documents": [doc.__dict__ for doc in documents[:20]],
        }

    def update_document_tags(self, document_id: int, tags: dict[str, str]) -> dict[str, Any]:
        self.init_schema()
        with self.connect() as connection:
            existing = self._fetch_one(connection, "SELECT id, metadata FROM lawagent_documents WHERE id = %s", (document_id,))
            if not existing:
                return {"error": "Document not found", "document_id": document_id}
            meta = existing.get("metadata") or {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except (json.JSONDecodeError, TypeError):
                    meta = {}
            for key in ("jurisdiction", "deal_stance", "deal_structure"):
                if key in tags:
                    meta[key] = tags[key]
            if self.backend == "postgres":
                self._execute(connection, "UPDATE lawagent_documents SET metadata = %s::jsonb WHERE id = %s", (json.dumps(meta), document_id))
            else:
                self._execute(connection, "UPDATE lawagent_documents SET metadata = %s WHERE id = %s", (json.dumps(meta), document_id))
            connection.commit()
        return {"status": "updated", "document_id": document_id, "tags": {k: meta.get(k, "") for k in ("jurisdiction", "deal_stance", "deal_structure")}}

    def _execute(self, connection, sql: str, params: tuple[Any, ...]):
        if self.backend == "sqlite":
            sql = sql.replace("%s", "?").replace("JSONB", "TEXT")
            sql = re.sub(r"::\w+", "", sql)
        return connection.execute(sql, params)

    def _fetch_one(self, connection, sql: str, params: tuple[Any, ...]):
        cursor = self._execute(connection, sql, params)
        row = cursor.fetchone()
        return dict(row) if row else None

    def _fetch_all(self, connection, sql: str, params: tuple[Any, ...]):
        cursor = self._execute(connection, sql, params)
        return [dict(row) for row in cursor.fetchall()]

    def _insert_document(
        self,
        connection,
        title: str,
        source_path: str,
        category: str,
        document_type: str,
        source_system: str,
        checksum: str,
        metadata: dict[str, Any],
    ) -> int:
        if self.backend == "postgres":
            row = connection.execute(
                """
                INSERT INTO lawagent_documents
                    (title, source_path, category, document_type, source_system, checksum, metadata)
                VALUES (%s, %s, %s, %s, %s, %s, %s::jsonb)
                RETURNING id
                """,
                (title, source_path, category, document_type, source_system, checksum, json.dumps(metadata)),
            ).fetchone()
            return int(row["id"])
        cursor = connection.execute(
            """
            INSERT INTO lawagent_documents
                (title, source_path, category, document_type, source_system, checksum, metadata, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (title, source_path, category, document_type, source_system, checksum, json.dumps(metadata), now_iso()),
        )
        return int(cursor.lastrowid)


_shared_db: CorpusDatabase | None = None
_db_lock = threading.Lock()


def get_db() -> CorpusDatabase:
    global _shared_db
    if _shared_db is None:
        with _db_lock:
            if _shared_db is None:
                _shared_db = CorpusDatabase()
    return _shared_db
