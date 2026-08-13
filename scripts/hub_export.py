"""Hub bake step — produces four download artifacts from a completed session.

Artifacts (all written to GCS at gcs_prefix/):
  1. redline.docx  — python-docx with <w:ins>/<w:del> track changes for accepted/edited
  2. clean.docx    — accepted/edited changes baked in, no track changes
  3. memo.docx     — issues memo with rationale, decisions log appendix
  4. register.json — structured array of all changes and user decisions

Footer on every Word artifact: "LawAgent demo — not legal advice."
"""

from __future__ import annotations

import io
import json
import logging
import os
from datetime import datetime, timezone
from typing import Any

logger = logging.getLogger(__name__)

_FOOTER = "LawAgent demo — not legal advice."


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def bake_session(
    session_id: str,
    draft_text: str,
    changes: list[dict[str, Any]],
    decisions: list[dict[str, Any]],
    posture: str = "neutral",
    gcs_prefix: str | None = None,
) -> dict[str, str]:
    """Produce all four artifacts for a session.

    Returns dict of {artifact_name: gcs_uri or local_path}.
    Writes to GCS when GCS_BUCKET env var is set; otherwise returns in-memory bytes
    under key suffix "_bytes" for testing.
    """
    accepted_changes = [
        c for c in changes if c.get("current_action") in ("accepted", "edited")
    ]
    rejected_changes = [
        c for c in changes if c.get("current_action") == "rejected"
    ]
    pending_changes = [
        c for c in changes if c.get("current_action") == "pending"
    ]

    redline_bytes = _build_redline_docx(draft_text, accepted_changes)
    clean_bytes = _build_clean_docx(draft_text, accepted_changes)
    memo_bytes = _build_memo_docx(changes, decisions, posture, session_id)
    register_bytes = _build_register_json(changes, decisions, session_id)

    artifacts: dict[str, Any] = {}

    bucket = os.environ.get("GCS_BUCKET", "")
    prefix = gcs_prefix or f"review-staging/{session_id}/"

    if bucket:
        artifacts["redline.docx"] = _upload_to_gcs(bucket, prefix + "redline.docx", redline_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        artifacts["clean.docx"] = _upload_to_gcs(bucket, prefix + "clean.docx", clean_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        artifacts["memo.docx"] = _upload_to_gcs(bucket, prefix + "memo.docx", memo_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        artifacts["register.json"] = _upload_to_gcs(bucket, prefix + "register.json", register_bytes, "application/json")
    else:
        artifacts["redline.docx_bytes"] = redline_bytes
        artifacts["clean.docx_bytes"] = clean_bytes
        artifacts["memo.docx_bytes"] = memo_bytes
        artifacts["register.json_bytes"] = register_bytes

    logger.info(
        "Bake complete for session %s: %d accepted, %d rejected, %d pending",
        session_id, len(accepted_changes), len(rejected_changes), len(pending_changes),
    )

    return artifacts


# ---------------------------------------------------------------------------
# Artifact builders
# ---------------------------------------------------------------------------

def _build_redline_docx(draft_text: str, accepted_changes: list[dict[str, Any]]) -> bytes:
    """Build a redlined .docx with track-change markup for accepted/edited changes."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    rev_id = 1  # OOXML requires unique w:id across all revision marks in the document

    paragraphs = draft_text.split("\n") if draft_text else ["[No draft text]"]
    for para_text in paragraphs:
        if not para_text.strip():
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()

        # Find if any accepted change covers this paragraph
        matching = [
            c for c in accepted_changes
            if c.get("original_text") and c["original_text"] in para_text
        ]

        if matching:
            # Apply all matching changes sequentially, tracking offset into remaining text
            remaining = para_text
            for change in matching:
                original = change["original_text"]
                replacement = change.get("current_text") or change.get("proposed_text", "")
                if original not in remaining:
                    continue
                idx = remaining.index(original)
                before = remaining[:idx]
                after = remaining[idx + len(original):]

                if before:
                    p.add_run(before)

                # Deletion: w:del → w:r → w:delText  (w:r wrapper required by OOXML schema)
                del_run = OxmlElement("w:del")
                del_run.set(qn("w:id"), str(rev_id)); rev_id += 1
                del_run.set(qn("w:author"), "LawAgent")
                del_run.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
                del_r = OxmlElement("w:r")
                del_text = OxmlElement("w:delText")
                del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                del_text.text = original
                del_r.append(del_text)
                del_run.append(del_r)
                p._element.append(del_run)

                # Insertion: w:ins → w:r → w:t
                ins_run = OxmlElement("w:ins")
                ins_run.set(qn("w:id"), str(rev_id)); rev_id += 1
                ins_run.set(qn("w:author"), "LawAgent")
                ins_run.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
                ins_r = OxmlElement("w:r")
                ins_text = OxmlElement("w:t")
                ins_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                ins_text.text = replacement
                ins_r.append(ins_text)
                ins_run.append(ins_r)
                p._element.append(ins_run)

                remaining = after  # continue from after this change

            if remaining:
                p.add_run(remaining)
        else:
            p.add_run(para_text)

    # New-section insertions: accepted changes with no original_text cannot
    # be spliced into a paragraph, so append them as tracked insertions at
    # the end of the document so Word's review pane sees them.
    for change in (c for c in accepted_changes if not c.get("original_text")):
        text = change.get("current_text") or change.get("proposed_text", "")
        if not text:
            continue
        p = doc.add_paragraph()
        anchor = change.get("clause_anchor", "")
        if anchor:
            note = p.add_run(f"[Proposed: {anchor}]  ")
            note.italic = True
        ins_elem = OxmlElement("w:ins")
        ins_elem.set(qn("w:id"), str(rev_id)); rev_id += 1
        ins_elem.set(qn("w:author"), "LawAgent")
        ins_elem.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
        ins_r = OxmlElement("w:r")
        ins_t = OxmlElement("w:t")
        ins_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        ins_t.text = text
        ins_r.append(ins_t)
        ins_elem.append(ins_r)
        p._element.append(ins_elem)

    _add_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_clean_docx(draft_text: str, accepted_changes: list[dict[str, Any]]) -> bytes:
    """Build a clean .docx with accepted/edited changes applied inline."""
    from docx import Document

    text = draft_text or ""
    for c in accepted_changes:
        original = c.get("original_text")
        replacement = c.get("current_text") or c.get("proposed_text", "")
        if original and original in text:
            text = text.replace(original, replacement)

    # New-section insertions: no original_text means the change is an addition,
    # not a replacement. Append at the end so the clean doc contains the clause.
    for c in accepted_changes:
        if not c.get("original_text"):
            addition = c.get("current_text") or c.get("proposed_text", "")
            if addition:
                text = text + "\n\n" + addition

    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    _add_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_memo_docx(
    changes: list[dict[str, Any]],
    decisions: list[dict[str, Any]],
    posture: str,
    session_id: str,
) -> bytes:
    """Build an issues memo with a decisions-log appendix."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    title = doc.add_heading("LawAgent Issues Memo", level=1)
    doc.add_paragraph(f"Session: {session_id}  |  Posture: {posture.upper()}  |  Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    severity_order = {"high": 0, "med": 1, "low": 2}
    sorted_changes = sorted(changes, key=lambda c: severity_order.get(c.get("severity", "med"), 1))

    redlines = [c for c in sorted_changes if c.get("kind") != "missing_clause"]
    missing = [c for c in sorted_changes if c.get("kind") == "missing_clause"]

    if redlines:
        doc.add_heading("Issues", level=2)
        for i, c in enumerate(redlines, 1):
            doc.add_heading(f"{i}. [{c.get('severity', 'med').upper()}] {c.get('clause_anchor', '')} — {c.get('category', '').replace('_', ' ').title()}", level=3)
            doc.add_paragraph(f"Why: {c.get('rationale', '')}")
            if c.get("source_ref"):
                doc.add_paragraph(f"Source: {c['source_ref']}")
            doc.add_paragraph(f"Proposed: {c.get('proposed_text', '')[:500]}")
            doc.add_paragraph(f"Decision: {c.get('current_action', 'pending').upper()}")
            doc.add_paragraph("")

    if missing:
        doc.add_heading("Missing Clauses", level=2)
        for i, c in enumerate(missing, 1):
            doc.add_heading(f"{i}. {c.get('category', '').replace('_', ' ').title()}", level=3)
            doc.add_paragraph(f"Why: {c.get('rationale', '')}")
            doc.add_paragraph(f"Proposed language: {c.get('proposed_text', '')[:800]}")
            doc.add_paragraph(f"Decision: {c.get('current_action', 'pending').upper()}")
            doc.add_paragraph("")

    # Decisions log appendix
    if decisions:
        doc.add_heading("Appendix: Decisions Log", level=2)
        for d in decisions:
            doc.add_paragraph(
                f"Change {d.get('change_id', '?')}: {d.get('action', '?').upper()} "
                f"at {d.get('ts', '?')}"
                + (f" — {d.get('edited_text', '')[:100]}" if d.get("edited_text") else "")
            )

    _add_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_register_json(
    changes: list[dict[str, Any]],
    decisions: list[dict[str, Any]],
    session_id: str,
) -> bytes:
    """Build the structured register.json artifact."""
    register = {
        "session_id": session_id,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "schema_version": "v1",
        "changes": [
            {
                "id": c.get("id", ""),
                "clause_anchor": c.get("clause_anchor", ""),
                "category": c.get("category", ""),
                "severity": c.get("severity", "med"),
                "kind": c.get("kind", "redline"),
                "original_text": c.get("original_text"),
                "proposed_text": c.get("proposed_text", ""),
                "current_action": c.get("current_action", "pending"),
                "current_text": c.get("current_text", ""),
                "rationale": c.get("rationale", ""),
                "source_ref": c.get("source_ref", ""),
            }
            for c in changes
        ],
        "decisions_log": decisions,
        "summary": {
            "total": len(changes),
            "accepted": sum(1 for c in changes if c.get("current_action") == "accepted"),
            "rejected": sum(1 for c in changes if c.get("current_action") == "rejected"),
            "edited": sum(1 for c in changes if c.get("current_action") == "edited"),
            "dismissed": sum(1 for c in changes if c.get("current_action") == "dismissed"),
            "pending": sum(1 for c in changes if c.get("current_action") == "pending"),
        },
    }
    return json.dumps(register, indent=2).encode()


def _add_footer(doc: Any) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    p = footer.paragraphs[0]
    p.clear()
    p.add_run(_FOOTER)


# ---------------------------------------------------------------------------
# GCS upload
# ---------------------------------------------------------------------------

def _upload_to_gcs(bucket: str, blob_name: str, data: bytes, content_type: str) -> str:
    try:
        from google.cloud import storage
        client = storage.Client()
        bucket_obj = client.bucket(bucket)
        blob = bucket_obj.blob(blob_name)
        blob.upload_from_string(data, content_type=content_type)
        return f"gs://{bucket}/{blob_name}"
    except Exception as exc:
        logger.error("GCS upload failed for %s/%s: %s", bucket, blob_name, exc)
        return ""
