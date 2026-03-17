"""Ingest legal documents into a local Chroma collection with rich metadata."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from collections import Counter, defaultdict
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Sequence

from docx import Document as DocxDocument
from docx.document import Document as DocxContainer
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

DEFAULT_DATA_DIR = Path("data")
DEFAULT_CHROMA_DIR = Path("chroma_db")
DEFAULT_COLLECTION = "ma_test"
DEFAULT_EMBED_MODEL = "nomic-embed-text"
DEFAULT_CHUNK_SIZE = 2000
DEFAULT_CHUNK_OVERLAP = 200
DEFAULT_PDF_MERGE_MIN_CHARS = 900
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

LEGAL_SEPARATORS = [
    r"\n\s*Section\s+\d+[A-Za-z]?\.",
    r"\n\s*Article\s+[IVXLC0-9]+",
    r"\n\s*\d+(?:\.\d+)*\.\s+",
    r"\n\s*\([a-z]\)\s+",
    r"\n\n",
    r"\n",
    r"\.\s+",
    r";\s+",
    r"\s+",
    "",
]

PAGE_NUMBER_RE = re.compile(r"^\s*(?:page\s+)?\d+\s*(?:of\s+\d+)?\s*$", flags=re.IGNORECASE)

JURISDICTION_HINTS = {
    "delaware": "Delaware",
    "newyork": "New York",
    "ny": "New York",
    "california": "California",
    "texas": "Texas",
    "uk": "UK",
    "unitedkingdom": "UK",
    "england": "UK",
    "eu": "EU",
    "europeanunion": "EU",
}

DOCUMENT_TYPE_HINTS = [
    ("sharepurchaseagreement", "Share Purchase Agreement"),
    ("assetpurchaseagreement", "Asset Purchase Agreement"),
    ("stockpurchaseagreement", "Stock Purchase Agreement"),
    ("mergeragreement", "Merger Agreement"),
    ("practice_note", "Practice Note"),
    ("practicenote", "Practice Note"),
    ("checklist", "Checklist"),
    ("clause", "Clause"),
    ("agreement", "Agreement"),
    ("table", "Comparison Table"),
]

PRACTICE_AREA_HINTS = [
    ("employment", "Employment"),
    ("ip", "IP"),
    ("privacy", "Privacy"),
    ("tax", "Tax"),
    ("regulatory", "Regulatory"),
    ("merger", "M&A"),
    ("acquisition", "M&A"),
    ("spa", "M&A"),
    ("apa", "M&A"),
]

CLAUSE_TYPE_HINTS = [
    ("indemn", "Indemnification"),
    ("changeofcontrol", "Change of Control"),
    ("assignment", "Assignment"),
    ("assignability", "Assignment"),
    ("materialadverse", "Material Adverse Change"),
    ("mac", "Material Adverse Change"),
    ("termination", "Termination"),
    ("governinglaw", "Governing Law"),
    ("noncompete", "Non-Compete"),
    ("nonsolicit", "Non-Solicit"),
]


def clear_db(chroma_dir: Path) -> None:
    if chroma_dir.exists():
        shutil.rmtree(chroma_dir)
    chroma_dir.mkdir(parents=True, exist_ok=True)


def normalize_token(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", text.lower())


def slugify(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")
    return slug or "document"


def extract_tokens(relative_path: Path) -> list[str]:
    cleaned = str(relative_path.with_suffix("")).lower().replace("&", " and ")
    return [token for token in re.split(r"[^a-z0-9]+", cleaned) if token]


def pick_first_match(
    haystack: str,
    tokens: set[str],
    hints: list[tuple[str, str]],
    default: str,
) -> str:
    for key, value in hints:
        normalized_key = normalize_token(key)
        if not normalized_key:
            continue
        if len(normalized_key) <= 3:
            if normalized_key in tokens:
                return value
            continue
        if normalized_key in haystack:
            return value
    return default


def infer_jurisdiction(haystack: str, tokens: set[str]) -> str:
    for key, value in JURISDICTION_HINTS.items():
        normalized_key = normalize_token(key)
        if len(normalized_key) <= 3:
            if normalized_key in tokens:
                return value
            continue
        if normalized_key in haystack:
            return value
    return "Unspecified"


def infer_source(relative_path: Path) -> str:
    parts = [part.lower() for part in relative_path.parts]
    if "westlaw" in parts or "practical_law" in parts or "practicallaw" in parts:
        return "Westlaw Practical Law"
    if parts:
        return parts[0].replace("_", " ").title()
    return "Unknown"


def normalize_metadata_value(value: Any) -> Any:
    if isinstance(value, (str, int, float, bool)):
        return value
    return str(value)


def load_metadata_overrides(metadata_json: Path | None) -> dict[str, dict[str, Any]]:
    if not metadata_json:
        return {}
    if not metadata_json.exists():
        raise FileNotFoundError(f"Metadata file not found: {metadata_json.resolve()}")

    payload = json.loads(metadata_json.read_text(encoding="utf-8-sig"))
    if not isinstance(payload, dict):
        raise ValueError("Metadata override file must be a JSON object.")

    normalized: dict[str, dict[str, Any]] = {}
    for key, value in payload.items():
        if isinstance(value, dict):
            normalized[str(key)] = {k: normalize_metadata_value(v) for k, v in value.items()}
    return normalized


def build_base_metadata(
    file_path: Path,
    data_dir: Path,
    metadata_overrides: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    relative_path = file_path.relative_to(data_dir)
    relative_posix = relative_path.as_posix()
    token_list = extract_tokens(relative_path)
    token_set = set(token_list)
    haystack = "".join(token_list)

    metadata = {
        "source": infer_source(relative_path),
        "source_file": relative_posix,
        "doc_id": slugify(str(relative_path.with_suffix(""))),
        "document_type": pick_first_match(
            haystack, token_set, DOCUMENT_TYPE_HINTS, "General Document"
        ),
        "jurisdiction": infer_jurisdiction(haystack, token_set),
        "practice_area": pick_first_match(haystack, token_set, PRACTICE_AREA_HINTS, "M&A"),
        "clause_type": pick_first_match(haystack, token_set, CLAUSE_TYPE_HINTS, "General"),
        "ingested_date": datetime.fromtimestamp(file_path.stat().st_mtime, tz=UTC).strftime(
            "%Y-%m-%d"
        ),
        "file_extension": file_path.suffix.lower(),
    }

    override = metadata_overrides.get(relative_posix) or metadata_overrides.get(file_path.name)
    if override:
        metadata.update({k: normalize_metadata_value(v) for k, v in override.items()})
    return metadata


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned_lines = []
    for line in text.split("\n"):
        squashed = re.sub(r"[ \t]+", " ", line).strip()
        if PAGE_NUMBER_RE.match(squashed):
            continue
        cleaned_lines.append(squashed)
    cleaned = "\n".join(cleaned_lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def normalize_inline_whitespace(text: str) -> str:
    return re.sub(r"[ \t]+", " ", text).strip()


def iter_docx_blocks(document: DocxContainer):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def collect_docx_noise_lines(docx_file: DocxDocument) -> set[str]:
    noise_lines: set[str] = set()
    for section in docx_file.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                text = normalize_inline_whitespace(paragraph.text)
                if text and len(text) <= 160:
                    noise_lines.add(text)
    return noise_lines


def render_docx_table(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [normalize_inline_whitespace(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)

    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]

    table_lines = [f"| {' | '.join(normalized_rows[0])} |"]
    if len(normalized_rows) > 1:
        table_lines.append(f"| {' | '.join(['---'] * width)} |")
        for row in normalized_rows[1:]:
            table_lines.append(f"| {' | '.join(row)} |")
    else:
        table_lines.append(f"| {' | '.join(['---'] * width)} |")

    return "\n".join(table_lines)


def trim_repeated_boundary_lines(text: str, noisy_lines: set[str]) -> str:
    lines = text.splitlines()
    while lines and lines[0].strip() in noisy_lines:
        lines.pop(0)
    while lines and lines[-1].strip() in noisy_lines:
        lines.pop()
    return "\n".join(lines)


def remove_repeated_pdf_headers_footers(docs: list[Document]) -> list[Document]:
    grouped: dict[str, list[Document]] = defaultdict(list)
    for doc in docs:
        if str(doc.metadata.get("file_extension", "")).lower() == ".pdf":
            grouped[str(doc.metadata.get("source_file", "unknown"))].append(doc)

    for _, page_docs in grouped.items():
        if len(page_docs) < 3:
            continue

        boundary_counts: Counter[str] = Counter()
        for page_doc in page_docs:
            lines = [line.strip() for line in page_doc.page_content.splitlines() if line.strip()]
            if not lines:
                continue
            if len(lines[0]) <= 120:
                boundary_counts[lines[0]] += 1
            if len(lines[-1]) <= 120:
                boundary_counts[lines[-1]] += 1

        threshold = max(2, int(len(page_docs) * 0.6))
        noisy_lines = {line for line, count in boundary_counts.items() if count >= threshold}
        if not noisy_lines:
            continue

        for page_doc in page_docs:
            page_doc.page_content = trim_repeated_boundary_lines(page_doc.page_content, noisy_lines)
    return docs


def merge_small_pdf_pages(docs: list[Document], min_chars: int) -> list[Document]:
    if min_chars <= 0:
        return docs

    grouped: dict[str, list[Document]] = defaultdict(list)
    non_pdf_docs: list[Document] = []

    for doc in docs:
        if str(doc.metadata.get("file_extension", "")).lower() == ".pdf":
            grouped[str(doc.metadata.get("source_file", "unknown"))].append(doc)
        else:
            non_pdf_docs.append(doc)

    merged_docs: list[Document] = []
    for _, page_docs in grouped.items():
        ordered = sorted(
            page_docs,
            key=lambda d: int(d.metadata.get("page", 10**9))
            if str(d.metadata.get("page", "")).isdigit()
            else 10**9,
        )

        buffer_text: list[str] = []
        page_start: int | None = None
        page_end: int | None = None
        base_metadata: dict[str, Any] | None = None

        def flush_buffer() -> None:
            nonlocal buffer_text, page_start, page_end, base_metadata
            if not buffer_text or base_metadata is None:
                return
            metadata = dict(base_metadata)
            metadata["page_start"] = page_start if page_start is not None else "n/a"
            metadata["page_end"] = page_end if page_end is not None else "n/a"
            if page_start is not None and page_end is not None:
                metadata["page"] = page_start if page_start == page_end else f"{page_start}-{page_end}"
            else:
                metadata["page"] = "n/a"
            merged_docs.append(Document(page_content="\n\n".join(buffer_text).strip(), metadata=metadata))
            buffer_text = []
            page_start = None
            page_end = None
            base_metadata = None

        for page_doc in ordered:
            text = page_doc.page_content.strip()
            if not text:
                continue

            page_value = page_doc.metadata.get("page")
            numeric_page = page_value if isinstance(page_value, int) else None
            if isinstance(page_value, str) and page_value.isdigit():
                numeric_page = int(page_value)

            if base_metadata is None:
                base_metadata = dict(page_doc.metadata)
                page_start = numeric_page

            page_marker = f"[PAGE {numeric_page}]\n{text}" if numeric_page is not None else text
            buffer_text.append(page_marker)
            if numeric_page is not None:
                page_end = numeric_page

            if sum(len(part) for part in buffer_text) >= min_chars:
                flush_buffer()

        flush_buffer()

    def page_sort_value(doc: Document) -> int:
        value = doc.metadata.get("page_start", doc.metadata.get("page", "n/a"))
        if isinstance(value, int):
            return value
        if isinstance(value, str):
            if value.isdigit():
                return int(value)
            if "-" in value:
                first = value.split("-", maxsplit=1)[0]
                if first.isdigit():
                    return int(first)
        return 10**9

    optimized_docs = non_pdf_docs + merged_docs
    optimized_docs.sort(
        key=lambda d: (
            str(d.metadata.get("source_file", "")),
            page_sort_value(d),
        )
    )
    return optimized_docs


def load_pdf(file_path: Path, data_dir: Path, metadata_overrides: dict[str, dict[str, Any]]) -> list[Document]:
    try:
        # Layout mode typically preserves legal section formatting better.
        loader = PyPDFLoader(str(file_path), extraction_mode="layout")
    except TypeError:
        loader = PyPDFLoader(str(file_path))
    docs = loader.load()
    base_metadata = build_base_metadata(file_path, data_dir, metadata_overrides)
    for doc in docs:
        doc.metadata.update(base_metadata)
        page = doc.metadata.get("page")
        if isinstance(page, int):
            doc.metadata["page"] = page + 1
    return docs


def load_docx(
    file_path: Path,
    data_dir: Path,
    metadata_overrides: dict[str, dict[str, Any]],
) -> list[Document]:
    docx_file = DocxDocument(str(file_path))
    noise_lines = collect_docx_noise_lines(docx_file)
    blocks: list[str] = []

    for block in iter_docx_blocks(docx_file):
        if isinstance(block, Paragraph):
            text = normalize_inline_whitespace(block.text)
            if not text or text in noise_lines:
                continue

            style_name = normalize_inline_whitespace(getattr(block.style, "name", "")).lower()
            if "heading" in style_name or "title" in style_name:
                blocks.append(f"Section {text}")
            elif "list" in style_name:
                blocks.append(f"- {text}")
            else:
                blocks.append(text)
        else:
            table_text = render_docx_table(block)
            if table_text:
                blocks.append(f"[TABLE]\n{table_text}")

    text = "\n\n".join(blocks).strip()
    if not text:
        return []
    metadata = build_base_metadata(file_path, data_dir, metadata_overrides)
    metadata["page"] = "n/a"
    metadata["has_tables"] = any(block.startswith("[TABLE]") for block in blocks)
    return [Document(page_content=text, metadata=metadata)]


def load_text_file(
    file_path: Path,
    data_dir: Path,
    metadata_overrides: dict[str, dict[str, Any]],
) -> list[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    metadata = build_base_metadata(file_path, data_dir, metadata_overrides)
    metadata["page"] = "n/a"
    return [Document(page_content=text, metadata=metadata)]


def detect_section_heading(text: str) -> str:
    preview = text[:1200]
    patterns = [
        r"(?im)^\s*((?:Section|Article)\s+[0-9IVXLC]+[^\n]{0,140})$",
        r"(?im)^\s*(\d+(?:\.\d+)*\.\s+[^\n]{0,140})$",
        r"(?im)^\s*([A-Z][A-Z \-/]{5,120})$",
    ]
    for pattern in patterns:
        match = re.search(pattern, preview)
        if match:
            return " ".join(match.group(1).split())[:160]

    first_line = next((line.strip() for line in preview.splitlines() if line.strip()), "")
    return first_line[:160] if first_line else "Unspecified"


def load_documents(
    data_dir: Path,
    metadata_overrides: dict[str, dict[str, Any]],
    skip_clean: bool,
    pdf_merge_min_chars: int,
) -> list[Document]:
    data_dir.mkdir(parents=True, exist_ok=True)
    files = sorted(path for path in data_dir.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS)
    if not files:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        print(f"No files found under {data_dir.resolve()} (supported: {supported})")
        return []

    docs: list[Document] = []
    for file_path in files:
        try:
            suffix = file_path.suffix.lower()
            if suffix == ".pdf":
                docs.extend(load_pdf(file_path, data_dir, metadata_overrides))
            elif suffix == ".docx":
                docs.extend(load_docx(file_path, data_dir, metadata_overrides))
            else:
                docs.extend(load_text_file(file_path, data_dir, metadata_overrides))
        except Exception as exc:  # pragma: no cover - runtime loader errors
            print(f"Skipped {file_path}: {exc}")

    docs = remove_repeated_pdf_headers_footers(docs)
    if not skip_clean:
        for doc in docs:
            doc.page_content = clean_text(doc.page_content)
        docs = [doc for doc in docs if doc.page_content.strip()]

    docs = merge_small_pdf_pages(docs, min_chars=pdf_merge_min_chars)

    print(f"Loaded {len(docs)} normalized document units from {len(files)} files.")
    return docs


def chunk_documents(docs: Sequence[Document], chunk_size: int, chunk_overlap: int) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=LEGAL_SEPARATORS,
        is_separator_regex=True,
        add_start_index=True,
    )
    chunks = splitter.split_documents(docs)

    chunk_counter_by_doc: Counter[str] = Counter()
    for chunk in chunks:
        doc_id = str(chunk.metadata.get("doc_id", "unknown"))
        chunk_counter_by_doc[doc_id] += 1
        chunk.metadata["chunk_index"] = chunk_counter_by_doc[doc_id]
        chunk.metadata["section_heading"] = detect_section_heading(chunk.page_content)

    doc_type_counts = Counter(str(chunk.metadata.get("document_type", "Unknown")) for chunk in chunks)
    print(f"Created {len(chunks)} chunks (size={chunk_size}, overlap={chunk_overlap}).")
    print("Chunk distribution by document_type:")
    for doc_type, count in doc_type_counts.most_common():
        print(f"  - {doc_type}: {count}")
    return chunks


def ingest(
    data_dir: Path,
    chroma_dir: Path,
    collection_name: str,
    embed_model: str,
    chunk_size: int,
    chunk_overlap: int,
    pdf_merge_min_chars: int,
    reset_db: bool,
    metadata_overrides: dict[str, dict[str, Any]],
    skip_clean: bool,
) -> int:
    if reset_db:
        clear_db(chroma_dir)
    else:
        chroma_dir.mkdir(parents=True, exist_ok=True)

    docs = load_documents(
        data_dir=data_dir,
        metadata_overrides=metadata_overrides,
        skip_clean=skip_clean,
        pdf_merge_min_chars=pdf_merge_min_chars,
    )
    if not docs:
        print("No documents were loaded; ingestion skipped.")
        return 1

    chunks = chunk_documents(docs, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    embeddings = OllamaEmbeddings(model=embed_model)
    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=str(chroma_dir),
        collection_name=collection_name,
    )
    print(
        "Ingestion complete.\n"
        f"Collection: {collection_name}\n"
        f"Persisted at: {chroma_dir.resolve()}"
    )
    return 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Ingest legal documents into Chroma.")
    parser.add_argument("--data-dir", type=Path, default=DEFAULT_DATA_DIR)
    parser.add_argument("--chroma-dir", type=Path, default=DEFAULT_CHROMA_DIR)
    parser.add_argument("--collection", default=DEFAULT_COLLECTION)
    parser.add_argument("--embed-model", default=DEFAULT_EMBED_MODEL)
    parser.add_argument("--chunk-size", type=int, default=DEFAULT_CHUNK_SIZE)
    parser.add_argument("--chunk-overlap", type=int, default=DEFAULT_CHUNK_OVERLAP)
    parser.add_argument(
        "--pdf-merge-min-chars",
        type=int,
        default=DEFAULT_PDF_MERGE_MIN_CHARS,
        help="Merge adjacent PDF pages until this many characters are accumulated.",
    )
    parser.add_argument(
        "--metadata-json",
        type=Path,
        default=None,
        help="Optional JSON file with per-document metadata overrides.",
    )
    parser.add_argument(
        "--skip-clean",
        action="store_true",
        help="Skip text cleanup and keep raw extracted text.",
    )
    parser.add_argument(
        "--no-reset",
        action="store_true",
        help="Do not delete existing chroma directory before ingesting.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.chunk_size <= 0:
        raise SystemExit("chunk-size must be > 0")
    if args.chunk_overlap < 0:
        raise SystemExit("chunk-overlap must be >= 0")
    if args.chunk_overlap >= args.chunk_size:
        raise SystemExit("chunk-overlap must be smaller than chunk-size")
    if args.pdf_merge_min_chars < 0:
        raise SystemExit("pdf-merge-min-chars must be >= 0")

    metadata_overrides = load_metadata_overrides(args.metadata_json)
    return ingest(
        data_dir=args.data_dir,
        chroma_dir=args.chroma_dir,
        collection_name=args.collection,
        embed_model=args.embed_model,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        pdf_merge_min_chars=args.pdf_merge_min_chars,
        reset_db=not args.no_reset,
        metadata_overrides=metadata_overrides,
        skip_clean=args.skip_clean,
    )


if __name__ == "__main__":
    sys.exit(main())
